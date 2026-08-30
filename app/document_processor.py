"""File upload validation — size, disguised-type, path-traversal defenses."""
import os
import re
import fitz
import docx
import magic

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024
ALLOWED_MIME_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}

def sanitize_filename(filename: str) -> str:
    filename = os.path.basename(filename)
    filename = re.sub(r"[^A-Za-z0-9._-]", "_", filename)
    return filename[:255]

def validate_and_extract(file_bytes: bytes, declared_filename: str) -> str:
    if len(file_bytes) == 0:
        raise ValueError("Empty file.")
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File exceeds {MAX_FILE_SIZE_BYTES // (1024*1024)}MB limit.")
    detected_mime = magic.from_buffer(file_bytes, mime=True)
    if detected_mime not in ALLOWED_MIME_TYPES:
        raise ValueError(f"Unsupported or disguised file type: {detected_mime}")
    sanitize_filename(declared_filename)
    if detected_mime == "application/pdf":
        return _extract_pdf(file_bytes)
    elif "wordprocessingml" in detected_mime:
        return _extract_docx(file_bytes)
    else:
        return file_bytes.decode("utf-8", errors="replace")[:200_000]

def _extract_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    if doc.page_count > 500:
        raise ValueError("PDF too large: max 500 pages.")
    text_parts = [page.get_text() for page in doc]
    doc.close()
    return "\n".join(text_parts)[:200_000]

def _extract_docx(file_bytes: bytes) -> str:
    import io
    document = docx.Document(io.BytesIO(file_bytes))
    if len(document.paragraphs) > 20_000:
        raise ValueError("DOCX too large: exceeds paragraph safety limit.")
    return "\n".join(p.text for p in document.paragraphs)[:200_000]